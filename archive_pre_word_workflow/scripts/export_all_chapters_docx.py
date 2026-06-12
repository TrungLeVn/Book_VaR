from __future__ import annotations

import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
CHAPTERS_DIR = PROJECT_ROOT / "chapters"
ARTIFACTS_DIR = PROJECT_ROOT / "artifacts"
BUILD_DIR = PROJECT_ROOT / ".docx_build"

QUARTO = Path(r"C:\Program Files\RStudio\resources\app\bin\quarto\bin\quarto.exe")
R_SCRIPT = Path(r"C:\Program Files\R\R-4.5.0\bin\Rscript.exe")


@dataclass(frozen=True)
class ChapterSpec:
    source: str
    chapter_no: int
    output: str


CHAPTER_SPECS = [
    ChapterSpec("ch01.qmd", 1, "chuong_1_sach_chuyen_khao.docx"),
    ChapterSpec("ch02.qmd", 2, "chuong_2_sach_chuyen_khao.docx"),
    ChapterSpec("ch03.qmd", 3, "chuong_3_sach_chuyen_khao.docx"),
    ChapterSpec("ch04.qmd", 4, "chuong_4_sach_chuyen_khao.docx"),
    ChapterSpec("ch05.qmd", 5, "chuong_5_sach_chuyen_khao.docx"),
]


QMD_FRONT_MATTER = """---
bibliography: ../references.bib
csl: ../apa.csl
prefer-html: false
format:
  docx:
    toc: false
    number-sections: true
    number-depth: 3
execute:
  echo: false
  warning: false
  message: false
crossref:
  fig-title: "Hình"
  tbl-title: "Bảng"
  sec-prefix: "Mục"
---

"""


CH01_REPLACEMENTS = {
    """library(tibble)
library(knitr)
library(kableExtra)

price_A_0 <- 50
price_A_1 <- 55
price_B_0 <- 200
price_B_1 <- 206

tbl_demo <- tibble(
  `Tài sản` = c("A", "B"),
  `Giá đầu kỳ` = c(price_A_0, price_B_0),
  `Giá cuối kỳ` = c(price_A_1, price_B_1),
  `Thay đổi giá` = c(price_A_1 - price_A_0, price_B_1 - price_B_0),
  `Tỷ lệ sinh lời (%)` = c(
    (price_A_1 / price_A_0 - 1) * 100,
    (price_B_1 / price_B_0 - 1) * 100
  )
)

tbl_demo |>
  kbl(
    digits = 2,
    align = "ccccc",
    booktabs = TRUE
  ) |>
  kable_styling(
    latex_options = "hold_position",
    full_width = FALSE,
    position = "center"
  )""": """library(tibble)
library(knitr)

price_A_0 <- 50
price_A_1 <- 55
price_B_0 <- 200
price_B_1 <- 206

tbl_demo <- tibble(
  `Tài sản` = c("A", "B"),
  `Giá đầu kỳ` = c(price_A_0, price_B_0),
  `Giá cuối kỳ` = c(price_A_1, price_B_1),
  `Thay đổi giá` = c(price_A_1 - price_A_0, price_B_1 - price_B_0),
  `Tỷ lệ sinh lời (%)` = c(
    (price_A_1 / price_A_0 - 1) * 100,
    (price_B_1 / price_B_0 - 1) * 100
  )
)

tbl_demo |>
  knitr::kable(
    digits = 2,
    align = "ccccc"
  )""",
    """library(tibble)
library(knitr)
library(kableExtra)

vol_tbl <- tribble(
  ~`Thước đo`, ~`Dữ liệu đầu vào`, ~`Ý tưởng chính`, ~`Ứng dụng phù hợp`,
  "Độ lệch chuẩn mẫu", "Chuỗi tỷ lệ sinh lời", "Đo mức phân tán trung bình của return trong toàn bộ mẫu", "Mô tả rủi ro tổng quát trong một giai đoạn",
  "Historical volatility", "Return quá khứ", "Ước lượng volatility từ dữ liệu lịch sử", "So sánh rủi ro giữa tài sản hoặc giữa giai đoạn",
  "Rolling volatility", "Return trong cửa sổ trượt", "Tính độ lệch chuẩn trên một số quan sát gần nhất", "Theo dõi biến động rủi ro theo thời gian",
  "Realized volatility", "Dữ liệu tần suất cao hoặc return trong ngày", "Tổng hợp biến động thực tế từ nhiều return ngắn hạn", "Đánh giá và so sánh dự báo volatility",
  "Implied volatility", "Giá quyền chọn", "Suy ra volatility hàm ý từ giá quyền chọn trên thị trường", "Đánh giá kỳ vọng thị trường về biến động tương lai",
  "Conditional volatility", "Mô hình ARCH/GARCH", "Ước lượng volatility có điều kiện theo thông tin quá khứ", "Dự báo rủi ro và tính VaR thay đổi theo thời gian"
)

vol_tbl |>
  kbl(booktabs = TRUE, align = "llll") |>
  kable_styling(
    latex_options = c("hold_position", "scale_down"),
    full_width = FALSE,
    position = "center"
  )""": """library(tibble)
library(knitr)

vol_tbl <- tribble(
  ~`Thước đo`, ~`Dữ liệu đầu vào`, ~`Ý tưởng chính`, ~`Ứng dụng phù hợp`,
  "Độ lệch chuẩn mẫu", "Chuỗi tỷ lệ sinh lời", "Đo mức phân tán trung bình của return trong toàn bộ mẫu", "Mô tả rủi ro tổng quát trong một giai đoạn",
  "Historical volatility", "Return quá khứ", "Ước lượng volatility từ dữ liệu lịch sử", "So sánh rủi ro giữa tài sản hoặc giữa giai đoạn",
  "Rolling volatility", "Return trong cửa sổ trượt", "Tính độ lệch chuẩn trên một số quan sát gần nhất", "Theo dõi biến động rủi ro theo thời gian",
  "Realized volatility", "Dữ liệu tần suất cao hoặc return trong ngày", "Tổng hợp biến động thực tế từ nhiều return ngắn hạn", "Đánh giá và so sánh dự báo volatility",
  "Implied volatility", "Giá quyền chọn", "Suy ra volatility hàm ý từ giá quyền chọn trên thị trường", "Đánh giá kỳ vọng thị trường về biến động tương lai",
  "Conditional volatility", "Mô hình ARCH/GARCH", "Ước lượng volatility có điều kiện theo thông tin quá khứ", "Dự báo rủi ro và tính VaR thay đổi theo thời gian"
)

vol_tbl |>
  knitr::kable(align = "llll")""",
    """library(tibble)
library(knitr)
library(kableExtra)

bridge_tbl <- tribble(
  ~`Đặc điểm quan sát được`, ~`Công cụ nhận diện`, ~`Hàm ý về volatility`, ~`Hàm ý đối với VaR`,
  "Return dao động quanh trung bình gần 0", "Đồ thị chuỗi return, thống kê mô tả", "Mean có thể đơn giản nhưng variance vẫn cần mô hình hóa", "VaR thường phụ thuộc nhiều vào volatility hơn mean",
  "Return không chuẩn, có đuôi dày", "Histogram, Q--Q plot, JB, AD, Lilliefors", "Cần xem xét phân phối sai số khác chuẩn", "VaR chuẩn có thể đánh giá thấp tail risk",
  "Return ít tự tương quan", "ACF/PACF và Ljung--Box trên return", "Mean return khó dự báo tuyến tính", "Mean model có thể đơn giản",
  "Squared return có tự tương quan", "Ljung--Box trên squared return hoặc absolute return", "Có dấu hiệu volatility clustering", "VaR nên thay đổi theo thời gian",
  "Có hiệu ứng ARCH", "ARCH-LM test", "Cần mô hình phương sai có điều kiện", "GARCH-based VaR có cơ sở thực nghiệm hơn"
)

bridge_tbl |>
  kbl(booktabs = TRUE, align = "llll") |>
  kable_styling(
    latex_options = c("hold_position", "scale_down"),
    full_width = FALSE,
    position = "center"
  )""": """library(tibble)
library(knitr)

bridge_tbl <- tribble(
  ~`Đặc điểm quan sát được`, ~`Công cụ nhận diện`, ~`Hàm ý về volatility`, ~`Hàm ý đối với VaR`,
  "Return dao động quanh trung bình gần 0", "Đồ thị chuỗi return, thống kê mô tả", "Mean có thể đơn giản nhưng variance vẫn cần mô hình hóa", "VaR thường phụ thuộc nhiều vào volatility hơn mean",
  "Return không chuẩn, có đuôi dày", "Histogram, Q--Q plot, JB, AD, Lilliefors", "Cần xem xét phân phối sai số khác chuẩn", "VaR chuẩn có thể đánh giá thấp tail risk",
  "Return ít tự tương quan", "ACF/PACF và Ljung--Box trên return", "Mean return khó dự báo tuyến tính", "Mean model có thể đơn giản",
  "Squared return có tự tương quan", "Ljung--Box trên squared return hoặc absolute return", "Có dấu hiệu volatility clustering", "VaR nên thay đổi theo thời gian",
  "Có hiệu ứng ARCH", "ARCH-LM test", "Cần mô hình phương sai có điều kiện", "GARCH-based VaR có cơ sở thực nghiệm hơn"
)

bridge_tbl |>
  knitr::kable(align = "llll")""",
}


def replace_exact(body: str, replacements: dict[str, str]) -> str:
    updated = body
    for old, new in replacements.items():
        if old not in updated:
            raise ValueError("Không tìm thấy đoạn cần thay thế trong chương 1.")
        updated = updated.replace(old, new)
    return updated


def sanitize_source(spec: ChapterSpec, source_text: str) -> str:
    body = source_text
    if spec.chapter_no == 1:
        body = replace_exact(body, CH01_REPLACEMENTS)
        body = body.replace("=======\n", "")
        body = body.replace("# Student-t được scale để có cùng độ lệch chuẩn với phân phối chuẩn\n", "")
    if spec.chapter_no == 5:
        body = body.replace(
            "expected = alpha * n(),",
            "expected = dplyr::first(alpha) * dplyr::n(),",
        )
        body = body.replace(
            'kupiec_LR = kupiec_test(exceed, alpha)["LRuc"],',
            'kupiec_LR = kupiec_test(exceed, dplyr::first(alpha))["LRuc"],',
        )
        body = body.replace(
            'kupiec_p = kupiec_test(exceed, alpha)["p_value"],',
            'kupiec_p = kupiec_test(exceed, dplyr::first(alpha))["p_value"],',
        )
        body = body.replace(
            'christoffersen_LRcc = christoffersen_test(exceed, alpha)["LRcc"],',
            'christoffersen_LRcc = christoffersen_test(exceed, dplyr::first(alpha))["LRcc"],',
        )
        body = body.replace(
            'christoffersen_p = christoffersen_test(exceed, alpha)["p_cc"],',
            'christoffersen_p = christoffersen_test(exceed, dplyr::first(alpha))["p_cc"],',
        )
        body = body.replace(
            'dq_stat = dq_test(exceed, VaR, alpha)["DQ"],',
            'dq_stat = dq_test(exceed, VaR, dplyr::first(alpha))["DQ"],',
        )
        body = body.replace(
            'dq_p = dq_test(exceed, VaR, alpha)["p_value"],',
            'dq_p = dq_test(exceed, VaR, dplyr::first(alpha))["p_value"],',
        )
    return body


def set_run_font(run, size: float, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_runs(paragraph, size: float, *, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size, bold=bold, italic=italic)


def set_paragraph_common(paragraph, *, align, before=3, after=3, line=1.15, first_line_cm: float | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Cm(first_line_cm) if first_line_cm is not None else None
    paragraph.alignment = align


def insert_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_run_font(run, 12)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def normalize_crossref_text(text: str) -> str:
    normalized = text.replace("Bảng Table\xa0", "Bảng ")
    normalized = normalized.replace("Hình Figure\xa0", "Hình ")
    normalized = normalized.replace("Bảng Table ", "Bảng ")
    normalized = normalized.replace("Hình Figure ", "Hình ")
    normalized = normalized.replace("Figure\xa0", "Hình ")
    normalized = normalized.replace("Figure ", "Hình ")
    normalized = normalized.replace("Table\xa0", "Bảng ")
    normalized = normalized.replace("Table ", "Bảng ")
    normalized = re.sub(r"Bảng \?@[\w\-]+", "Bảng dưới đây", normalized)
    normalized = re.sub(r"Hình \?@[\w\-]+", "Hình dưới đây", normalized)
    normalized = re.sub(r";\s*[A-Za-z][A-Za-z0-9_]*\d{4}\?", "", normalized)
    normalized = re.sub(r"\(\s*([A-Za-z][A-Za-z0-9_]*\d{4}\?)\s*;\s*", "(", normalized)
    normalized = re.sub(r"\(\s*([A-Za-z][A-Za-z0-9_]*\d{4}\?)\s*\)", "", normalized)
    normalized = normalized.replace(" ;", ";")
    normalized = normalized.replace("(;", "(")
    normalized = normalized.replace(";)", ")")
    normalized = re.sub(r"\(\s*\)", "", normalized)
    return normalized


def uppercase_heading_title(text: str) -> str:
    match = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)$", text.strip())
    if not match:
        return text.upper()
    return f"{match.group(1)} {match.group(2).upper()}"


def configure_layout(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)
        section.start_type = WD_SECTION_START.CONTINUOUS

        sect_pr = section._sectPr
        if sect_pr.find(qn("w:mirrorMargins")) is None:
            sect_pr.append(OxmlElement("w:mirrorMargins"))

        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_page_number(paragraph)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.space_before = Pt(3)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, align in (
        ("Heading 1", 14, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 13, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 13, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)


def style_tables(document: Document) -> None:
    for table in document.tables:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    set_paragraph_common(
                        paragraph,
                        align=WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                        before=0,
                        after=0,
                        line=1.0,
                        first_line_cm=0,
                    )
                    style_runs(paragraph, 12, bold=(row_index == 0))


def style_paragraphs(document: Document, chapter_no: int) -> None:
    first_heading_done = False
    in_references = False

    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""

        if not text:
            continue

        if style_name in {"Title", "Source Code"}:
            delete_paragraph(paragraph)
            continue

        if style_name == "Heading 1" and not first_heading_done:
            title_text = re.sub(r"^\d+\.?\s*", "", text).strip().upper()
            chapter_paragraph = paragraph.insert_paragraph_before(f"Chương {chapter_no}")
            set_paragraph_common(
                chapter_paragraph,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                before=6,
                after=3,
                line=1.15,
                first_line_cm=0,
            )
            style_runs(chapter_paragraph, 14, bold=True, italic=False)

            paragraph.clear()
            paragraph.add_run(title_text)
            set_paragraph_common(
                paragraph,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                before=3,
                after=6,
                line=1.15,
                first_line_cm=0,
            )
            style_runs(paragraph, 14, bold=True, italic=False)
            first_heading_done = True
            continue

        if style_name == "Heading 1":
            demoted_text = re.sub(r"^\d+\.?\s*", "", text).strip()
            paragraph.style = document.styles["Normal"]
            paragraph.clear()
            paragraph.add_run(demoted_text)
            text = demoted_text

        if style_name == "Heading 2":
            paragraph.clear()
            paragraph.add_run(uppercase_heading_title(text))
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=1.15, first_line_cm=0)
            style_runs(paragraph, 13, bold=True, italic=False)
            continue

        if style_name == "Heading 3":
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=1.15, first_line_cm=0)
            style_runs(paragraph, 13, bold=True, italic=False)
            continue

        if style_name == "Heading 4":
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3, line=1.15, first_line_cm=0)
            style_runs(paragraph, 13, bold=False, italic=True)
            continue

        normalized_text = normalize_crossref_text(text)
        if normalized_text != text:
            paragraph.clear()
            paragraph.add_run(normalized_text)
            text = normalized_text

        if text in {"References", "Tài liệu tham khảo"}:
            paragraph.clear()
            paragraph.add_run("TÀI LIỆU THAM KHẢO")
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6, line=1.15, first_line_cm=0)
            style_runs(paragraph, 13, bold=True, italic=False)
            in_references = True
            continue

        if text.startswith("Bảng "):
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3, line=1.0, first_line_cm=0)
            style_runs(paragraph, 12, bold=True, italic=False)
            continue

        if text.startswith("Hình "):
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3, line=1.0, first_line_cm=0)
            style_runs(paragraph, 12, bold=True, italic=False)
            continue

        if text.startswith("Nguồn:"):
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT, before=3, after=3, line=1.0, first_line_cm=0)
            style_runs(paragraph, 12, bold=False, italic=True)
            continue

        if style_name == "Bibliography":
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=3, line=1.15, first_line_cm=-0.8)
            paragraph.paragraph_format.left_indent = Cm(0.8)
            style_runs(paragraph, 12, bold=False, italic=False)
            in_references = True
            continue

        if in_references:
            set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=3, line=1.15, first_line_cm=-0.8)
            paragraph.paragraph_format.left_indent = Cm(0.8)
            style_runs(paragraph, 12, bold=False, italic=False)
            continue

        set_paragraph_common(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=3, after=3, line=1.15, first_line_cm=1)
        style_runs(paragraph, 13, bold=None, italic=None)


def build_temp_qmd(spec: ChapterSpec) -> Path:
    source_path = CHAPTERS_DIR / spec.source
    source_text = source_path.read_text(encoding="utf-8")
    sanitized = sanitize_source(spec, source_text)

    BUILD_DIR.mkdir(exist_ok=True)
    temp_qmd = BUILD_DIR / f"{source_path.stem}_docx_tmp.qmd"
    temp_qmd.write_text(QMD_FRONT_MATTER + sanitized, encoding="utf-8")
    return temp_qmd


def render_raw_docx(temp_qmd: Path, raw_docx: Path) -> None:
    env = os.environ.copy()
    env["LOCALAPPDATA"] = str((PROJECT_ROOT / ".quarto-local").resolve())
    env["QUARTO_R"] = str(R_SCRIPT)

    raw_docx.parent.mkdir(parents=True, exist_ok=True)
    local_output = PROJECT_ROOT / raw_docx.name
    if local_output.exists():
        local_output.unlink()

    command = [
        str(QUARTO),
        "render",
        str(temp_qmd.relative_to(PROJECT_ROOT)),
        "--to",
        "docx",
        "--output",
        local_output.name,
    ]
    subprocess.run(command, cwd=PROJECT_ROOT, env=env, check=True)
    if not local_output.exists():
        alt_output = temp_qmd.with_name(raw_docx.name)
        if alt_output.exists():
            local_output = alt_output
        else:
            raise FileNotFoundError(f"Không tìm thấy file DOCX thô sau khi render: {raw_docx.name}")

    shutil.move(str(local_output), str(raw_docx))


def postprocess_docx(raw_docx: Path, final_docx: Path, chapter_no: int) -> None:
    document = Document(raw_docx)
    configure_layout(document)
    configure_styles(document)
    style_paragraphs(document, chapter_no)
    style_tables(document)
    document.save(final_docx)


def structural_audit(docx_path: Path) -> list[str]:
    document = Document(docx_path)
    findings: list[str] = []

    for bad_token in ("?@", "Figure", "Table", "======="):
        if any(bad_token in p.text for p in document.paragraphs):
            findings.append(f"token:{bad_token}")

    if any(p.style.name == "Source Code" for p in document.paragraphs):
        findings.append("source-code-visible")

    return findings


def export_one(spec: ChapterSpec) -> tuple[Path, list[str]]:
    temp_qmd = build_temp_qmd(spec)
    raw_docx = BUILD_DIR / f"{Path(spec.output).stem}_raw.docx"
    final_docx = ARTIFACTS_DIR / spec.output

    render_raw_docx(temp_qmd, raw_docx)
    postprocess_docx(raw_docx, final_docx, spec.chapter_no)
    findings = structural_audit(final_docx)
    return final_docx, findings


def clean_build_artifacts() -> None:
    for target in (BUILD_DIR, PROJECT_ROOT / ".quarto-local"):
        if target.exists():
            shutil.rmtree(target, ignore_errors=True)


def main() -> None:
    ARTIFACTS_DIR.mkdir(exist_ok=True)
    results: list[tuple[str, list[str]]] = []
    failures: list[tuple[str, str]] = []
    try:
        for spec in CHAPTER_SPECS:
            try:
                final_docx, findings = export_one(spec)
                results.append((str(final_docx), findings))
            except Exception as exc:
                failures.append((spec.source, str(exc)))
    finally:
        clean_build_artifacts()

    for path, findings in results:
        if findings:
            print(f"{path} | WARN | {', '.join(findings)}")
        else:
            print(f"{path} | OK")

    for source, error in failures:
        print(f"{source} | ERROR | {error}")


if __name__ == "__main__":
    main()
